#!/usr/bin/env python3
"""Gera o documento Word (ABNT) descrevendo o projeto A3."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

FONTE = "Times New Roman"

doc = Document()

# ── Configuração de página: A4, margens ABNT ──────────────────────────────
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(3.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(2.0)

# ── Estilo base: Times New Roman 12, espaçamento simples ──────────────────
normal = doc.styles["Normal"]
normal.font.name = normal.font.name
normal.font.name = FONTE
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONTE)
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
pf.space_before = Pt(0)
pf.space_after = Pt(0)


def set_run(run, bold=False, italic=False, size=12):
    run.font.name = FONTE
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), FONTE)


def add_par(texto="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
            italic=False, size=12, first_line=None, space_after=0, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(before)
    if first_line is not None:
        p.paragraph_format.first_line_indent = Cm(first_line)
    if texto:
        run = p.add_run(texto)
        set_run(run, bold=bold, italic=italic, size=size)
    return p


def add_heading(texto):
    return add_par(texto, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, before=2)


# ── Título ────────────────────────────────────────────────────────────────
add_par(
    "Retrato Falado e Reconhecimento Facial: um sistema de composição "
    "gráfica e identificação de suspeitos baseado em LBPH",
    align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=6,
)

# ── Palavras-chave ──────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Palavras-chave: ")
set_run(r, bold=True)
r = p.add_run(
    "Reconhecimento facial; Retrato falado; Visão computacional; "
    "LBPH; Processamento de imagens."
)
set_run(r)

# ── Introdução ──────────────────────────────────────────────────────────────
add_heading("1 Introdução")
add_par(
    "O retrato falado é uma técnica tradicionalmente empregada na investigação "
    "criminal para reconstruir, a partir do relato de testemunhas, a aparência de "
    "um suspeito. Sua produção manual é, contudo, lenta e dependente da habilidade "
    "do desenhista, e o resultado raramente é confrontado de forma automática com "
    "bases de identificação. Este trabalho aborda esse problema com um sistema que "
    "monta o retrato falado a partir de componentes faciais e o compara com uma base "
    "de rostos. O objetivo é desenvolver uma aplicação que integre "
    "composição gráfica por camadas e reconhecimento facial, retornando os suspeitos "
    "mais semelhantes ao retrato montado.",
    first_line=1.25,
)

# ── Métodos ───────────────────────────────────────────────────────────────
add_heading("2 Métodos")
add_par(
    "O sistema foi implementado em Python (OpenCV, Pillow, NumPy e MediaPipe), com "
    "interface web em Flask, em um pipeline de quatro etapas. (i) Geração de assets: "
    "de fotografias reais, detectam-se marcos faciais, alinha-se cada rosto e "
    "recortam-se faixas transparentes para cabelo, sobrancelhas, olhos, nariz, boca e "
    "formato do rosto. (ii) Composição: o usuário escolhe um componente por categoria "
    "e o compositor sobrepõe as camadas (alpha compositing) em uma tela de 500x600 px, "
    "aplicando uma transformação afim que adapta cada componente à geometria do rosto. "
    "(iii) Pré-processamento: a imagem é convertida para cinza, a face é localizada por "
    "Haar Cascade e normalizada para 200x200 px. (iv) Reconhecimento: extrai-se o "
    "histograma de padrões binários locais (LBP) e compara-se à base por distância "
    "qui-quadrado, retornando os cinco mais próximos. A base de 20 sujeitos é treinada "
    "com o reconhecedor LBPH do OpenCV, com aumento de dados para tolerar variações "
    "entre o retrato montado e a face original.",
    first_line=1.25,
)

# ── Resultados e Discussão ───────────────────────────────────────────────────
add_heading("3 Resultados e Discussão")
add_par(
    "O sistema produz em tempo real um retrato falado coerente e o submete ao "
    "reconhecimento, exibindo os cinco melhores candidatos com suas distâncias e "
    "similaridades normalizadas. Espera-se que, quando o retrato é montado com "
    "componentes de um mesmo sujeito da base, este figure entre os primeiros "
    "colocados. Adotou-se um reconhecimento por região, no qual cada componente é "
    "comparado à mesma região de cada sujeito e as distâncias são somadas; essa "
    "abordagem mostrou-se mais robusta que o histograma holístico, pois evita que o "
    "fundo branco de retratos parciais domine a comparação. Como limitação, o "
    "desempenho depende do alinhamento e do tamanho reduzido da base, fatores que "
    "tendem a melhorar com mais sujeitos e fotografias padronizadas.",
    first_line=1.25,
)

# ── Considerações Finais ─────────────────────────────────────────────────────
add_heading("4 Considerações Finais")
add_par(
    "O trabalho demonstrou a viabilidade de integrar composição gráfica por camadas "
    "e reconhecimento facial em uma única aplicação interativa, cumprindo o objetivo "
    "proposto. As técnicas de visão computacional empregadas (LBP/LBPH, Haar Cascade "
    "e alinhamento por marcos faciais) mostraram-se adequadas ao contexto acadêmico. "
    "Como trabalhos futuros, sugere-se ampliar a base e incorporar reconhecimento "
    "baseado em aprendizado profundo.",
    first_line=1.25,
)

# ── Referências ──────────────────────────────────────────────────────────────
add_heading("Referências")


def add_ref(texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.widow_control = False
    run = p.add_run(texto)
    set_run(run)
    return p


add_ref(
    "AHONEN, T.; HADID, A.; PIETIKÄINEN, M. Face description with local binary "
    "patterns: application to face recognition. IEEE Transactions on Pattern "
    "Analysis and Machine Intelligence, v. 28, n. 12, p. 2037-2041, 2006."
)
add_ref(
    "BRADSKI, G. The OpenCV Library. Dr. Dobb's Journal of Software Tools, 2000."
)
add_ref(
    "VIOLA, P.; JONES, M. Rapid object detection using a boosted cascade of simple "
    "features. In: IEEE CVPR, 2001. Anais [...]. Kauai: IEEE, 2001. p. 511-518."
)

caminho = "output/A3_retrato_falado_reconhecimento_facial.docx"
import os
os.makedirs("output", exist_ok=True)
doc.save(caminho)
print(f"Documento salvo em: {caminho}")
